import os
import re
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

def parse_puml(filepath):
    with open(filepath, 'r', encoding='utf-8') as f:
        content = f.read()

    # Extract actors: actor "Quản trị viên\n(Admin)" as Admin
    # or actor Admin
    actors = {}
    for match in re.finditer(r'actor\s+(?:"([^"]+)"|(\S+))(?:\s+as\s+(\S+))?', content, re.IGNORECASE):
        name = match.group(1) or match.group(2)
        name = name.replace('\\n', ' ').replace('\n', ' ')
        alias = match.group(3) or name
        actors[alias] = name

    # Default actor if none found explicitly
    if not actors:
        actors['User'] = 'Người dùng'

    # Extract usecases: usecase "Tên" as UC_Alias
    usecases = {}
    for match in re.finditer(r'usecase\s+(?:"([^"]+)"|(\S+))(?:\s+as\s+(\S+))?', content, re.IGNORECASE):
        name = match.group(1) or match.group(2)
        name = name.replace('\\n', ' ').replace('\n', ' ')
        alias = match.group(3) or name
        usecases[alias] = name

    # Extract relations
    # We mainly care about MainUC and included UC, but any UC mapped to an actor is a root.
    # Actually, simpler: find MainUC if it exists.
    # A lot of diagrams have `usecase "..." as MainUC`.
    main_uc = None
    if 'MainUC' in usecases:
        main_uc = usecases['MainUC']
    
    # Or just document all usecases that are NOT MainUC (or document all of them)
    # The user asked for "phân rã usecase" so the included use cases are the specific ones.
    sub_ucs = []
    # Find includes: MainUC ..> UC_View : <<include>>
    for match in re.finditer(r'(\S+)\s+(?:\.\.>|-->)\s+(\S+)\s*:\s*<<include>>', content, re.IGNORECASE):
        parent, child = match.group(1), match.group(2)
        if child in usecases and child not in [u['alias'] for u in sub_ucs]:
            sub_ucs.append({'alias': child, 'name': usecases[child]})
            
    # If no includes found, maybe it's just a direct actor --> UC
    if not sub_ucs:
        for alias, name in usecases.items():
            if alias != 'MainUC':
                sub_ucs.append({'alias': alias, 'name': name})
                
    # If still empty, use MainUC
    if not sub_ucs and 'MainUC' in usecases:
        sub_ucs.append({'alias': 'MainUC', 'name': usecases['MainUC']})
        
    # Get primary actor
    primary_actor = list(actors.values())[0] if actors else "Người dùng"
    
    return {
        'main_title': main_uc or "Đặc tả Use Case",
        'actor': primary_actor,
        'sub_ucs': sub_ucs
    }

def generate_text_for_uc(uc_name, actor):
    uc_lower = uc_name.lower()
    
    # Default values
    muc_dich = f"Cho phép {actor} thực hiện chức năng {uc_name}."
    tien_dieu_kien = f"{actor} đã truy cập vào hệ thống."
    hau_dieu_kien = f"Hệ thống hoàn tất xử lý yêu cầu {uc_name}."
    luong_chinh = [
        f"1. {actor} chọn chức năng '{uc_name}' trên giao diện.",
        "2. Hệ thống hiển thị màn hình tương ứng.",
        f"3. {actor} thao tác và xác nhận yêu cầu.",
        "4. Hệ thống xử lý thông tin.",
        "5. Hệ thống hiển thị thông báo thành công."
    ]
    luong_phu = [
        "A1 - Lỗi kết nối: Hệ thống thông báo lỗi mạng và yêu cầu thử lại.",
        "A2 - Hủy bỏ thao tác: Người dùng chọn hủy, hệ thống quay lại màn hình trước."
    ]

    # Heuristics based on name
    if "thêm" in uc_lower or "tạo" in uc_lower or "đăng" in uc_lower and "nhập" not in uc_lower and "ký" not in uc_lower:
        muc_dich = f"Cho phép {actor} tạo mới dữ liệu ({uc_name.replace('Thêm', '').replace('Tạo', '').strip()}) vào hệ thống."
        tien_dieu_kien = f"{actor} đã đăng nhập và có quyền truy cập chức năng này."
        hau_dieu_kien = "Dữ liệu mới được lưu thành công vào cơ sở dữ liệu."
        luong_chinh = [
            f"1. {actor} chọn chức năng '{uc_name}'.",
            "2. Hệ thống hiển thị biểu mẫu (form) nhập liệu.",
            f"3. {actor} điền đầy đủ các thông tin yêu cầu và nhấn nút Lưu/Xác nhận.",
            "4. Hệ thống kiểm tra tính hợp lệ của dữ liệu (validation).",
            "5. Hệ thống lưu dữ liệu vào cơ sở dữ liệu.",
            "6. Hệ thống hiển thị thông báo thành công và làm mới giao diện."
        ]
        luong_phu.append("A3 - Dữ liệu không hợp lệ: Hệ thống bôi đỏ các trường sai, hiển thị thông báo và yêu cầu nhập lại.")
        
    elif "sửa" in uc_lower or "cập nhật" in uc_lower:
        muc_dich = f"Cho phép {actor} thay đổi thông tin đối tượng đã có."
        tien_dieu_kien = f"{actor} đang ở giao diện chi tiết hoặc danh sách đối tượng."
        hau_dieu_kien = "Dữ liệu mới được cập nhật vào cơ sở dữ liệu."
        luong_chinh = [
            f"1. {actor} chọn đối tượng cần sửa và nhấn nút 'Chỉnh sửa'.",
            "2. Hệ thống tải dữ liệu hiện tại và hiển thị lên biểu mẫu.",
            f"3. {actor} thay đổi thông tin và nhấn nút 'Lưu'.",
            "4. Hệ thống kiểm tra tính hợp lệ của thông tin mới.",
            "5. Hệ thống cập nhật cơ sở dữ liệu.",
            "6. Hệ thống hiển thị thông báo cập nhật thành công."
        ]
        luong_phu.append("A3 - Dữ liệu trống: Nếu các trường bắt buộc bị xóa trắng, hệ thống từ chối lưu và yêu cầu nhập lại.")
        
    elif "xóa" in uc_lower or "từ chối" in uc_lower:
        muc_dich = f"Cho phép {actor} loại bỏ hoặc hủy bỏ đối tượng khỏi hệ thống."
        tien_dieu_kien = f"{actor} có quyền xóa và đang xem đối tượng cần xóa."
        hau_dieu_kien = "Dữ liệu bị xóa khỏi hệ thống (xóa cứng hoặc mềm)."
        luong_chinh = [
            f"1. {actor} chọn đối tượng và nhấn nút 'Xóa'.",
            "2. Hệ thống hiển thị hộp thoại cảnh báo và yêu cầu xác nhận.",
            f"3. {actor} chọn 'Đồng ý' (hoặc 'Xác nhận').",
            "4. Hệ thống thực hiện xóa dữ liệu.",
            "5. Hệ thống thông báo thành công và cập nhật lại danh sách hiển thị."
        ]
        
    elif "xem" in uc_lower or "danh sách" in uc_lower or "tra cứu" in uc_lower or "thống kê" in uc_lower:
        muc_dich = f"Cho phép {actor} tra cứu và theo dõi thông tin liên quan đến {uc_name.replace('Xem', '').strip()}."
        tien_dieu_kien = f"{actor} đã đăng nhập thành công."
        hau_dieu_kien = "Hệ thống hiển thị dữ liệu chính xác và đầy đủ."
        luong_chinh = [
            f"1. {actor} truy cập vào trang/chức năng '{uc_name}'.",
            "2. Hệ thống gửi yêu cầu tải dữ liệu từ máy chủ.",
            "3. Hệ thống truy xuất cơ sở dữ liệu (áp dụng phân trang, bộ lọc nếu có).",
            "4. Hệ thống hiển thị danh sách dữ liệu lên giao diện."
        ]
        luong_phu.append("A3 - Không có dữ liệu: Nếu chưa có dữ liệu nào, hệ thống hiển thị thông báo 'Không tìm thấy dữ liệu'.")

    elif "đăng nhập" in uc_lower:
        muc_dich = "Xác thực danh tính người dùng để cho phép truy cập vào các chức năng bảo mật."
        tien_dieu_kien = "Người dùng chưa đăng nhập, đang ở màn hình Đăng nhập."
        hau_dieu_kien = "Hệ thống cấp Token/Session, người dùng đăng nhập thành công."
        luong_chinh = [
            "1. Người dùng nhập tên tài khoản (hoặc Email) và Mật khẩu.",
            "2. Người dùng nhấn nút 'Đăng nhập'.",
            "3. Hệ thống gửi thông tin đến máy chủ để xác thực.",
            "4. Hệ thống kiểm tra thông tin đối chiếu với cơ sở dữ liệu.",
            "5. Hệ thống lưu phiên làm việc và chuyển hướng vào trang chủ."
        ]
        luong_phu = [
            "A1 - Sai tài khoản/mật khẩu: Hệ thống thông báo 'Tài khoản hoặc mật khẩu không chính xác'.",
            "A2 - Tài khoản bị khóa: Hệ thống báo lỗi 'Tài khoản của bạn đang bị khóa'."
        ]

    elif "đăng ký" in uc_lower:
        muc_dich = "Cho phép người dùng mới tạo tài khoản để sử dụng hệ thống."
        tien_dieu_kien = "Người dùng truy cập vào trang Đăng ký."
        hau_dieu_kien = "Tài khoản mới được tạo và lưu vào hệ thống."
        luong_chinh = [
            "1. Người dùng nhập các thông tin bắt buộc: Tên, Email, Mật khẩu, Xác nhận mật khẩu.",
            "2. Người dùng nhấn 'Đăng ký'.",
            "3. Hệ thống kiểm tra tính hợp lệ của định dạng Email và độ mạnh Mật khẩu.",
            "4. Hệ thống kiểm tra Email đã tồn tại hay chưa.",
            "5. Hệ thống tạo tài khoản mới và có thể gửi email xác thực.",
            "6. Hệ thống hiển thị thông báo đăng ký thành công và chuyển đến trang Đăng nhập."
        ]
        luong_phu.append("A3 - Email đã tồn tại: Hệ thống thông báo 'Email đã được sử dụng'.")

    elif "bình luận" in uc_lower:
        muc_dich = f"Cho phép {actor} tương tác bằng cách {uc_name}."
        tien_dieu_kien = f"{actor} đang xem chi tiết nội dung (truyện/chương)."
        hau_dieu_kien = "Nội dung tương tác được ghi nhận và hiển thị công khai."
        luong_chinh = [
            f"1. {actor} nhập nội dung bình luận vào ô nhập liệu.",
            f"2. {actor} nhấn nút 'Gửi' (hoặc 'Bình luận').",
            "3. Hệ thống kiểm tra nội dung (không được trống, lọc từ ngữ vi phạm).",
            "4. Hệ thống lưu bình luận vào cơ sở dữ liệu.",
            "5. Hệ thống cập nhật hiển thị bình luận mới ngay trên giao diện."
        ]

    elif "thanh toán" in uc_lower or "mua" in uc_lower:
        muc_dich = f"Cho phép {actor} thực hiện giao dịch tài chính để mua gói dịch vụ."
        tien_dieu_kien = f"{actor} đã chọn gói dịch vụ và tiến hành thanh toán."
        hau_dieu_kien = "Giao dịch thành công, tài khoản được cập nhật quyền lợi tương ứng."
        luong_chinh = [
            f"1. {actor} chọn phương thức thanh toán.",
            "2. Hệ thống chuyển hướng sang cổng thanh toán (Payment Gateway).",
            f"3. {actor} thực hiện thao tác chuyển khoản hoặc quẹt thẻ.",
            "4. Cổng thanh toán trả về kết quả giao dịch cho hệ thống.",
            "5. Hệ thống xác nhận thanh toán thành công và cập nhật quyền lợi (VD: VIP).",
            "6. Hệ thống hiển thị thông báo giao dịch hoàn tất."
        ]
        luong_phu.append("A3 - Thanh toán thất bại/Hủy: Người dùng hủy thanh toán hoặc số dư không đủ, hệ thống thông báo 'Giao dịch thất bại'.")

    elif "phân quyền" in uc_lower or "cấp role" in uc_lower:
        muc_dich = f"Cho phép {actor} quản lý và thay đổi vai trò (Role) của người dùng khác."
        tien_dieu_kien = f"{actor} là Admin và đang ở trang quản trị người dùng."
        hau_dieu_kien = "Vai trò người dùng được cập nhật trong hệ thống."
        luong_chinh = [
            f"1. {actor} chọn một người dùng và nhấn 'Phân quyền'.",
            "2. Hệ thống hiển thị danh sách các vai trò (Role) khả dụng.",
            f"3. {actor} chọn một vai trò mới và nhấn 'Cập nhật'.",
            "4. Hệ thống yêu cầu xác nhận thao tác.",
            "5. Hệ thống lưu thông tin quyền mới vào cơ sở dữ liệu.",
            "6. Hệ thống thông báo thành công và tải lại bảng dữ liệu."
        ]

    elif "khóa" in uc_lower:
        muc_dich = f"Cho phép {actor} thay đổi trạng thái hoạt động của đối tượng (Khóa/Mở khóa)."
        hau_dieu_kien = "Trạng thái đối tượng được cập nhật."
        luong_chinh = [
            f"1. {actor} chọn đối tượng và nhấn nút 'Khóa/Mở khóa'.",
            "2. Hệ thống hiển thị hộp thoại xác nhận.",
            f"3. {actor} nhấn 'Đồng ý'.",
            "4. Hệ thống cập nhật trạng thái trong cơ sở dữ liệu.",
            "5. Hệ thống hiển thị thông báo thành công."
        ]

    return {
        'muc_dich': muc_dich,
        'tien_dieu_kien': tien_dieu_kien,
        'hau_dieu_kien': hau_dieu_kien,
        'luong_chinh': luong_chinh,
        'luong_phu': luong_phu
    }

def create_word_document(parsed_data, filename):
    doc = Document()
    
    # Title
    title = doc.add_heading(f"Đặc tả Use Case: {parsed_data['main_title']}", 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Intro
    doc.add_paragraph("Tài liệu này cung cấp mô tả chi tiết cho các Use Case thuộc phân hệ này, nhằm phục vụ quá trình phát triển và kiểm thử hệ thống.")
    
    for idx, uc in enumerate(parsed_data['sub_ucs'], 1):
        uc_name = uc['name']
        actor = parsed_data['actor']
        
        doc.add_heading(f"{idx}. Use Case: {uc_name}", level=1)
        
        spec = generate_text_for_uc(uc_name, actor)
        
        # Info block
        doc.add_paragraph().add_run(f"Tác nhân (Actor): ").bold = True
        doc.paragraphs[-1].add_run(actor)
        
        doc.add_paragraph().add_run(f"Mục đích: ").bold = True
        doc.paragraphs[-1].add_run(spec['muc_dich'])
        
        doc.add_paragraph().add_run(f"Tiền điều kiện: ").bold = True
        doc.paragraphs[-1].add_run(spec['tien_dieu_kien'])
        
        doc.add_paragraph().add_run(f"Hậu điều kiện: ").bold = True
        doc.paragraphs[-1].add_run(spec['hau_dieu_kien'])
        
        # Main flow
        doc.add_heading("Luồng sự kiện chính (Main Flow):", level=2)
        for step in spec['luong_chinh']:
            doc.add_paragraph(step, style='List Number')
            
        # Alt flow
        doc.add_heading("Luồng sự kiện thay thế / Ngoại lệ (Alternative Flows):", level=2)
        for alt in spec['luong_phu']:
            doc.add_paragraph(alt, style='List Bullet')
            
        doc.add_paragraph("\n") # spacing

    doc.save(filename)

def main():
    directory = "."
    for file in os.listdir(directory):
        if file.endswith(".puml") and file.startswith("phan_ra"):
            filepath = os.path.join(directory, file)
            print(f"Processing {file}...")
            try:
                parsed = parse_puml(filepath)
                # print(parsed)
                out_name = file.replace(".puml", ".docx")
                create_word_document(parsed, out_name)
                print(f" -> Created {out_name}")
            except Exception as e:
                print(f"Error processing {file}: {e}")

if __name__ == "__main__":
    main()
